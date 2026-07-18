#!/usr/bin/env python3
"""
Reporte Ejecutivo Contable — Sucesion
Genera reporte_ejecutivo_sucesion.docx con analisis completo de
ingresos, egresos, balances, graficas y desglose por año y mes.
"""

import io
import os
import re
import sys
import unicodedata
from datetime import datetime

import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mtick
import requests

try:
    from google.oauth2.service_account import Credentials
    from google.auth.transport.requests import Request as GoogleRequest
    import gspread
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Dependencias faltantes. Ejecuta:  pip install -r requirements.txt")
    sys.exit(1)

# ── Config ─────────────────────────────────────────────────────────────────────
SPREADSHEET_ID   = "1YTvNIui0kBSMWRs6mMrZvbM690ILMl_W"
SHEET_NAME_MOV   = "movimientos"
CREDENTIALS_FILE = "credentials.json"
OUTPUT_FILE      = "reporte_ejecutivo_sucesion.docx"
SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets.readonly",
    "https://www.googleapis.com/auth/drive.readonly",
]

MESES_CORTO = ["Ene","Feb","Mar","Abr","May","Jun",
               "Jul","Ago","Sep","Oct","Nov","Dic"]
MESES_LARGO = ["Enero","Febrero","Marzo","Abril","Mayo","Junio",
               "Julio","Agosto","Septiembre","Octubre","Noviembre","Diciembre"]

COLOR_ING  = "#2E86C1"
COLOR_EGR  = "#E74C3C"
COLOR_BAL  = "#27AE60"
COLOR_DARK = "#1F4E79"

# ── Auth / descarga ────────────────────────────────────────────────────────────
def get_credentials():
    if not os.path.exists(CREDENTIALS_FILE):
        print(f"ERROR: No se encontro '{CREDENTIALS_FILE}'.")
        sys.exit(1)
    return Credentials.from_service_account_file(CREDENTIALS_FILE, scopes=SCOPES)

def download_movimientos() -> pd.DataFrame:
    creds = get_credentials()
    try:
        client = gspread.authorize(creds)
        ws  = client.open_by_key(SPREADSHEET_ID).worksheet(SHEET_NAME_MOV)
        df  = pd.DataFrame(ws.get_all_records()).fillna("")
        return df
    except gspread.exceptions.APIError as e:
        if "not supported" not in str(e) and "Office file" not in str(e):
            raise
    creds.refresh(GoogleRequest())
    url  = f"https://docs.google.com/spreadsheets/d/{SPREADSHEET_ID}/export?format=xlsx"
    resp = requests.get(url, headers={"Authorization": f"Bearer {creds.token}"}, timeout=60)
    resp.raise_for_status()
    all_sheets = pd.read_excel(io.BytesIO(resp.content), sheet_name=None, dtype=str)
    for name, df in all_sheets.items():
        if SHEET_NAME_MOV.lower() in name.lower():
            return df.fillna("")
    return list(all_sheets.values())[0].fillna("")

# ── Helpers de datos ───────────────────────────────────────────────────────────
def clean_numeric(s):
    return pd.to_numeric(s.astype(str).str.replace(r"[$,\s%]","",regex=True), errors="coerce")

def parse_dates(s):
    return pd.to_datetime(s, errors="coerce")

MESES_MAP = {
    "enero":1,"febrero":2,"marzo":3,"abril":4,"mayo":5,"junio":6,
    "julio":7,"agosto":8,"septiembre":9,"octubre":10,"noviembre":11,"diciembre":12,
    "ene":1,"feb":2,"mar":3,"abr":4,"may":5,"jun":6,
    "jul":7,"ago":8,"sep":9,"oct":10,"nov":11,"dic":12,
}

def parse_mes_col(series: pd.Series) -> pd.Series:
    result = []
    for val in series:
        v = str(val).strip().lower()
        try:
            n = int(float(v))
            result.append(n if 1 <= n <= 12 else None)
            continue
        except (ValueError, TypeError):
            pass
        num = MESES_MAP.get(v)
        if num is None:
            for name, n in MESES_MAP.items():
                if name in v:
                    num = n
                    break
        result.append(num)
    return pd.array(result, dtype="Int64")

def detect_cols(df):
    def find(kws):
        for c in df.columns:
            if any(k in c.lower() for k in kws): return c
        return None
    def find_exact(names):
        for c in df.columns:
            norm = unicodedata.normalize("NFC", c.strip().lower())
            if norm in names: return c
        for c in df.columns:
            norm = unicodedata.normalize("NFC", c.strip().lower())
            for n in names:
                if re.match(rf"^{re.escape(n)}\b", norm): return c
        return None
    return (find({"fecha","date"}),
            find({"ingreso","entrada","cobro","haber"}),
            find({"egreso","salida","gasto","debe"}),
            find({"concepto","descripcion","detalle"}),
            find({"registro","cuenta","rubro","categoria"}),
            find_exact({"año","anio","year","ejercicio"}),
            find_exact({"mes","month","periodo"}),
            find({"emisor","proveedor","razon social","razon_social","rfc emisor"}))

def money(v):
    return f"${v:,.2f}" if pd.notna(v) else "$0.00"

# ── Graficas ───────────────────────────────────────────────────────────────────
def _buf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    return buf

def chart_barras_mensual(monthly: pd.DataFrame, año=None) -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(10, 4))
    x  = range(len(monthly))
    w  = 0.35
    ax.bar([i - w/2 for i in x], monthly["ing"], w, label="Ingresos",
           color=COLOR_ING, alpha=0.85)
    ax.bar([i + w/2 for i in x], monthly["egr"], w, label="Egresos",
           color=COLOR_EGR, alpha=0.85)
    ax.set_xticks(list(x))
    ax.set_xticklabels(monthly["mes"], rotation=45, ha="right", fontsize=8)
    ax.yaxis.set_major_formatter(mtick.FuncFormatter(lambda v,_: f"${v:,.0f}"))
    ax.legend(fontsize=9)
    ax.set_title(f"Ingresos vs Egresos{' — ' + str(int(año)) if año else ''}", fontsize=11, fontweight="bold")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _buf(fig)

def chart_balance_linea(monthly: pd.DataFrame, año=None) -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(10, 3.5))
    bal = monthly["ing"] - monthly["egr"]
    colors = [COLOR_BAL if v >= 0 else COLOR_EGR for v in bal]
    ax.bar(range(len(bal)), bal, color=colors, alpha=0.8)
    ax.axhline(0, color="#555", linewidth=0.8, linestyle="--")
    ax.set_xticks(range(len(monthly)))
    ax.set_xticklabels(monthly["mes"], rotation=45, ha="right", fontsize=8)
    ax.yaxis.set_major_formatter(mtick.FuncFormatter(lambda v,_: f"${v:,.0f}"))
    ax.set_title(f"Balance Mensual{' — ' + str(int(año)) if año else ''}", fontsize=11, fontweight="bold")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _buf(fig)

def chart_pie_registro(reg_df: pd.DataFrame, col_label: str) -> io.BytesIO:
    top = reg_df.nlargest(8, "ing")
    fig, ax = plt.subplots(figsize=(9, 5))
    wedges, _, autotexts = ax.pie(
        top["ing"].clip(lower=0),
        autopct="%1.1f%%",
        startangle=140,
        pctdistance=0.75,
    )
    for t in autotexts: t.set_fontsize(7)
    ax.legend(wedges, top[col_label].tolist(),
              loc="center left", bbox_to_anchor=(1, 0.5), fontsize=7)
    ax.set_title("Distribucion de Ingresos por Registro", fontsize=10, fontweight="bold")
    fig.tight_layout()
    return _buf(fig)

def chart_anual_barras(anual: pd.DataFrame) -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(9, 4))
    x = range(len(anual))
    w = 0.28
    ax.bar([i - w for i in x], anual["ing"], w, label="Ingresos", color=COLOR_ING, alpha=0.85)
    ax.bar([i      for i in x], anual["egr"], w, label="Egresos",  color=COLOR_EGR, alpha=0.85)
    ax.bar([i + w  for i in x], anual["ing"] - anual["egr"], w,
           label="Balance", color=COLOR_BAL, alpha=0.85)
    ax.set_xticks(list(x))
    ax.set_xticklabels([str(int(a)) for a in anual["año"]], fontsize=9)
    ax.yaxis.set_major_formatter(mtick.FuncFormatter(lambda v,_: f"${v:,.0f}"))
    ax.legend(fontsize=9)
    ax.set_title("Resumen Anual — Ingresos / Egresos / Balance", fontsize=11, fontweight="bold")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _buf(fig)

# ── Helpers Word ───────────────────────────────────────────────────────────────
def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p   = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold      = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = _rgb(color)

def add_heading(doc, text, level=1, color=COLOR_DARK):
    p   = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = _rgb(color)
    return p

def add_kpi_table(doc, kpis: list):
    """kpis = [(label, value), ...]"""
    n   = len(kpis)
    tbl = doc.add_table(rows=2, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style     = "Table Grid"
    for ci, (lbl, val) in enumerate(kpis):
        hc = tbl.cell(0, ci)
        vc = tbl.cell(1, ci)
        set_cell_bg(hc, "1F4E79")
        set_cell_bg(vc, "D6E4F0")
        cell_text(hc, lbl,  bold=True, size=9,  color="#FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(vc, val,  bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in tbl.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return tbl

def add_data_table(doc, headers: list, rows: list, col_widths=None):
    n_cols = len(headers)
    tbl    = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style     = "Table Grid"

    # 7pt para tablas anchas (5+ columnas), 8pt para las demás
    fsize = 7 if n_cols >= 5 else 8

    # Encabezados
    for ci, hdr in enumerate(headers):
        cell = tbl.cell(0, ci)
        set_cell_bg(cell, "2E86C1")
        cell_text(cell, hdr, bold=True, size=fsize, color="#FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Datos
    for ri, row_data in enumerate(rows):
        first = str(row_data[0]).upper()
        is_total    = first.startswith("TOTAL")
        is_subtotal = first.startswith("SUBTOTAL")
        bg = "2E86C1" if is_subtotal else ("D6E4F0" if is_total else ("F2F2F2" if ri % 2 == 0 else "FFFFFF"))
        txt_color = "#FFFFFF" if is_subtotal else None
        for ci, val in enumerate(row_data):
            cell = tbl.cell(ri+1, ci)
            set_cell_bg(cell, bg)
            align = WD_ALIGN_PARAGRAPH.RIGHT if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            cell_text(cell, val, bold=(is_total or is_subtotal), size=fsize,
                      color=txt_color, align=align)

    # Anchos de columna — siempre se aplican; auto-calcula si no se especifican
    if not col_widths:
        _total = 16.0
        _first = min(5.0, round(_total * 0.35, 2))
        _rest  = round((_total - _first) / max(n_cols - 1, 1), 2)
        col_widths = [_first] + [_rest] * (n_cols - 1)
    for ci, w in enumerate(col_widths):
        for row in tbl.rows:
            row.cells[ci].width = Cm(w)
    return tbl

def add_image(doc, buf: io.BytesIO, width=Inches(6)):
    doc.add_picture(buf, width=width)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

if getattr(sys, 'frozen', False):
    _BASE_DIR = os.path.dirname(os.path.abspath(sys.executable))
else:
    _BASE_DIR = os.path.dirname(os.path.abspath(__file__))
XLSX_DASHBOARD = os.path.join(_BASE_DIR, "dashboard_sucesion.xlsx")

def insertar_hoja_control(doc, sheet_name):
    """Lee la hoja pivot del Excel y la inserta como tabla Word formateada.
    2025 → Jul-Dic | 2026 → Ene-Jun | formato $peso | anchos proporcionales."""
    import openpyxl as _opxl

    if not os.path.exists(XLSX_DASHBOARD):
        doc.add_paragraph("[dashboard_sucesion.xlsx no encontrado — ejecute dashboard.py primero]")
        return

    wb = _opxl.load_workbook(XLSX_DASHBOARD, read_only=True)
    if sheet_name not in wb.sheetnames:
        wb.close()
        doc.add_paragraph(f"[Hoja '{sheet_name}' no encontrada en el Excel]")
        return

    ws = wb[sheet_name]
    all_rows = [list(r[1:]) for r in ws.iter_rows(values_only=True)
                if not all(v is None for v in r)]
    wb.close()
    if not all_rows:
        return

    # Formato peso: $4,800.00 (sin espacio para que quepa en columnas angostas)
    def fmt_peso(v):
        if v is None: return "—"
        if isinstance(v, (int, float)):
            if v == 0: return "—"
            return f"${v:,.2f}"
        return str(v).strip() or "—"

    # Columnas a mostrar por año (índices en la fila original de 14 cols)
    # Estructura: [entity(0), Ene(1)..Dic(12), Total(13)]
    KEEP = {
        '2025': [0, 7, 8, 9, 10, 11, 12, 13],   # entity + Jul-Dic + Total
        '2026': [0, 1, 2, 3, 4, 5, 6, 13],      # entity + Ene-Jun + Total
    }

    def filtrar(row, indices):
        n = len(row)
        return [row[i] if 0 <= i < n else None for i in indices]

    sheet_title = str(all_rows[0][0] or '').strip()
    max_cols = max(len(r) for r in all_rows)

    # Parsear bloques por año
    blocks, gran_total_val = [], None
    cur_year = cur_hdrs = None
    cur_data = []

    for row in all_rows:
        while len(row) < max_cols: row.append(None)
        first = str(row[0] or '').strip()
        rest_none = all(v is None for v in row[1:])

        if first.upper() == 'GRAN TOTAL':
            if cur_year and cur_hdrs:
                blocks.append((cur_year, cur_hdrs, cur_data))
                cur_year = cur_hdrs = None; cur_data = []
            nums = [v for v in row[1:] if isinstance(v, (int, float))]
            gran_total_val = nums[-1] if nums else None
            continue
        if rest_none and re.search(r'\d{4}', first):
            if cur_year and cur_hdrs:
                blocks.append((cur_year, cur_hdrs, cur_data))
            cur_year = re.search(r'\d{4}', first).group()
            cur_hdrs = None; cur_data = []
            continue
        if len(row) > 1 and str(row[1] or '').strip() == 'Ene':
            cur_hdrs = [str(row[0] or '').strip()] + [str(v or '').strip() for v in row[1:]]
            continue
        if rest_none:
            continue
        if cur_year is not None and cur_hdrs is not None:
            cur_data.append(row)

    if cur_year and cur_hdrs and cur_data:
        blocks.append((cur_year, cur_hdrs, cur_data))
    if not blocks:
        doc.add_paragraph("[Sin datos en la hoja de control]")
        return

    # n_cols después del filtro (igual para todos los años: 8 columnas)
    first_keep = KEEP.get(blocks[0][0], list(range(len(blocks[0][1]))))
    n_cols = len(first_keep)

    # Construir lista plana: (tipo, fila_ya_filtrada_y_formateada)
    flat = [('title', [sheet_title])]
    for year_label, headers, rows_data in blocks:
        keep = KEEP.get(year_label, list(range(len(headers))))
        filt_hdrs = filtrar(headers, keep)
        flat.append(('year_hdr', [f'AÑO {year_label}']))
        flat.append(('col_hdr', [str(v or '') for v in filt_hdrs]))
        data_alt = 0
        for row in rows_data:
            first_str = str(row[0] or '').strip()
            is_total  = first_str.upper().startswith('TOTAL')
            filt = filtrar(row, keep)
            fmt_row = [str(filt[0] or '').strip()] + [fmt_peso(v) for v in filt[1:]]
            rtype = 'total' if is_total else ('data_even' if data_alt % 2 == 0 else 'data_odd')
            flat.append((rtype, fmt_row))
            data_alt = 0 if is_total else data_alt + 1

    if gran_total_val is not None:
        flat.append(('gran_total', [f'GRAN TOTAL   {fmt_peso(gran_total_val)}']))

    BG = {
        'title':      '1F3864', 'year_hdr':  '1F4E79',
        'col_hdr':    '2E86C1', 'total':     '2E86C1',
        'gran_total': '1F3864', 'data_even': 'FFFFFF', 'data_odd': 'F2F2F2',
    }
    DARK  = {'title', 'year_hdr', 'col_hdr', 'total', 'gran_total'}
    MERGE = {'title', 'year_hdr', 'gran_total'}

    tbl = doc.add_table(rows=len(flat), cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Anchos proporcionales al contenido dentro de 16 cm disponibles
    entity_cm = 3.5
    total_cm  = 2.2
    month_cm  = round((16.0 - entity_cm - total_cm) / max(n_cols - 2, 1), 3)
    col_w = [entity_cm] + [month_cm] * (n_cols - 2) + [total_cm]

    for ri, (rtype, row_data) in enumerate(flat):
        bg    = BG[rtype]
        bold  = rtype in DARK
        size  = 9 if rtype in ('title', 'gran_total') else (8 if rtype in ('year_hdr', 'col_hdr') else 7)
        color = 'FFFFFF' if rtype in DARK else '333333'

        if rtype in MERGE:
            merged = tbl.cell(ri, 0).merge(tbl.cell(ri, n_cols - 1))
            set_cell_bg(merged, bg)
            cell_text(merged, row_data[0], bold=True, size=size,
                      color=f'#{color}', align=WD_ALIGN_PARAGRAPH.CENTER)
        elif rtype == 'col_hdr':
            for ci in range(n_cols):
                cell = tbl.cell(ri, ci)
                set_cell_bg(cell, bg)
                val  = row_data[ci] if ci < len(row_data) else ''
                aln  = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
                cell_text(cell, str(val), bold=True, size=size, color=f'#{color}', align=aln)
        else:
            for ci in range(n_cols):
                cell = tbl.cell(ri, ci)
                set_cell_bg(cell, bg)
                val  = row_data[ci] if ci < len(row_data) else '—'
                aln  = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.RIGHT
                cell_text(cell, str(val), bold=bold, size=size, color=f'#{color}', align=aln)

    # Fijar anchos en TODAS las filas no fusionadas para que Word los respete
    for ri, (rtype, _) in enumerate(flat):
        if rtype not in MERGE:
            for ci in range(n_cols):
                tbl.cell(ri, ci).width = Cm(col_w[ci])


def insertar_detalle_registro(doc):
    """Lee la hoja 'Detalle Registro' del Excel e inserta una sección por cada
    tipo de registro con tabla Inmueble|Año|Mes|Ingreso|Egreso|Concepto."""
    import openpyxl as _opxl

    SHEET = "Detalle Registro"
    if not os.path.exists(XLSX_DASHBOARD):
        doc.add_paragraph("[dashboard_sucesion.xlsx no encontrado — ejecute dashboard.py primero]")
        return

    wb = _opxl.load_workbook(XLSX_DASHBOARD, read_only=True)
    if SHEET not in wb.sheetnames:
        wb.close()
        doc.add_paragraph(f"[Hoja '{SHEET}' no encontrada en el Excel]")
        return

    ws = wb[SHEET]
    # Leer cols B-G (índices 1-6), saltar col A (indent)
    raw = [list(r[1:7]) for r in ws.iter_rows(values_only=True)
           if not all(v is None for v in r)]
    wb.close()

    def fmt_peso(v):
        if v is None or v == '': return ""
        if isinstance(v, (int, float)):
            if v == 0: return "—"
            return f"$ {v:,.2f}"
        return str(v).strip()

    SKIP_WORDS = ('DETALLE DE MOVIMIENTOS', 'GENERADO:')

    # Parsear secciones: cada sección = (registro, [(inm,año,mes,ing,egr,conc)], sub_ing, sub_egr)
    sections, gran_ing, gran_egr = [], 0.0, 0.0
    cur_reg, cur_rows = None, []

    for row in raw:
        c0 = str(row[0] or '').strip()
        rest_blank = all(v is None or str(v).strip() == '' for v in row[1:])

        # Encabezado de sección: c0 tiene texto y el resto son None (celda fusionada)
        if rest_blank and c0 and not any(w in c0.upper() for w in SKIP_WORDS) \
                and not c0.upper().startswith('SUBTOTAL') and c0.upper() != 'GRAN TOTAL':
            if cur_reg is not None:          # guardar sección anterior si quedó sin subtotal
                sections.append((cur_reg, cur_rows, 0.0, 0.0))
            cur_reg, cur_rows = c0, []
            continue

        if c0.upper() == 'GRAN TOTAL':
            gran_ing = float(row[3]) if isinstance(row[3], (int, float)) else 0.0
            gran_egr = float(row[4]) if isinstance(row[4], (int, float)) else 0.0
            break

        if c0.upper().startswith('SUBTOTAL'):
            sub_ing = float(row[3]) if isinstance(row[3], (int, float)) else 0.0
            sub_egr = float(row[4]) if isinstance(row[4], (int, float)) else 0.0
            if cur_reg is not None:
                sections.append((cur_reg, cur_rows, sub_ing, sub_egr))
            cur_reg, cur_rows = None, []
            continue

        # Fila de encabezados de columna → saltar
        if c0 == 'Inmueble':
            continue

        # Fila de datos
        if cur_reg is not None and c0:
            cur_rows.append(row)

    if not sections:
        doc.add_paragraph("[Sin datos en la hoja Detalle Registro]")
        return

    COL_W = [4.0, 1.3, 1.5, 2.5, 2.5, 4.2]   # Inmueble|Año|Mes|Ingreso|Egreso|Concepto (total 16.0 cm)
    HDR   = ["Inmueble", "Año", "Mes", "Ingreso", "Egreso", "Concepto"]

    for reg_name, rows_data, sub_ing, sub_egr in sections:
        add_heading(doc, reg_name, level=2)
        word_rows = []
        for row in rows_data:
            inm, año, mes, ing, egr, conc = row
            word_rows.append([
                str(inm  or ''),
                str(año  or ''),
                str(mes  or ''),
                fmt_peso(ing),
                fmt_peso(egr),
                str(conc or ''),
            ])
        if sub_ing or sub_egr:
            word_rows.append(["SUBTOTAL", "", "", fmt_peso(sub_ing), fmt_peso(sub_egr), ""])
        add_data_table(doc, HDR, word_rows, col_widths=COL_W)
        doc.add_paragraph()

    if gran_ing or gran_egr:
        add_kpi_table(doc, [
            ("Gran Total Ingresos", fmt_peso(gran_ing)),
            ("Gran Total Egresos",  fmt_peso(gran_egr)),
        ])
        doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# GENERADOR DEL REPORTE
# ══════════════════════════════════════════════════════════════════════════════
def generate_report(df: pd.DataFrame, output: str):
    fecha_col, ing_col, egr_col, conc_col, reg_col, año_col, mes_col, emisor_col = detect_cols(df)

    # ── Preparar datos base ────────────────────────────────────────────────────
    tmp = df.copy()
    tmp["_ing"] = clean_numeric(tmp[ing_col]) if ing_col else 0
    tmp["_egr"] = clean_numeric(tmp[egr_col]) if egr_col else 0
    tmp["_bal"] = tmp["_ing"].fillna(0) - tmp["_egr"].fillna(0)

    if año_col:
        _s = pd.to_numeric(tmp[año_col].astype(str).str.strip(), errors="coerce").round(0)
        _na = _s.isna().to_numpy()
        tmp["_año"] = pd.arrays.IntegerArray(_s.fillna(0).to_numpy().astype("int64"), _na)
    elif fecha_col:
        tmp["_año"] = parse_dates(tmp[fecha_col]).dt.year.astype("Int64")
    else:
        tmp["_año"] = pd.array([pd.NA] * len(tmp), dtype="Int64")

    if mes_col:
        tmp["_mes"] = parse_mes_col(tmp[mes_col])
    elif fecha_col:
        tmp["_mes"] = parse_dates(tmp[fecha_col]).dt.month.astype("Int64")
    else:
        tmp["_mes"] = pd.array([pd.NA] * len(tmp), dtype="Int64")

    total_ing = tmp["_ing"].sum()
    total_egr = tmp["_egr"].sum()
    total_bal = total_ing - total_egr

    años = sorted(tmp["_año"].dropna().unique().astype(int))
    periodo = f"{años[0]} – {años[-1]}" if años else "—"

    # ── Mensual global ─────────────────────────────────────────────────────────
    def monthly_df(subset):
        g = (subset.groupby("_mes")
             .agg(ing=("_ing","sum"), egr=("_egr","sum"))
             .reindex(range(1,13), fill_value=0)
             .reset_index()
             .rename(columns={"_mes":"mes_num"}))
        g["mes"] = g["mes_num"].apply(lambda m: MESES_CORTO[m-1])
        return g

    # ── Anual ──────────────────────────────────────────────────────────────────
    anual = (tmp.groupby("_año")
             .agg(ing=("_ing","sum"), egr=("_egr","sum"))
             .reset_index()
             .rename(columns={"_año":"año"}))
    anual["bal"] = anual["ing"] - anual["egr"]

    # ── Por Registro ───────────────────────────────────────────────────────────
    reg_df = None
    if reg_col:
        reg_df = (tmp.groupby(reg_col)
                  .agg(ing=("_ing","sum"), egr=("_egr","sum"))
                  .reset_index())
        reg_df["bal"] = reg_df["ing"] - reg_df["egr"]
        reg_df = reg_df.sort_values("ing", ascending=False)

    # ══════════════════════════════════════════════════════════════════════════
    print("  Construyendo documento Word...")
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── PORTADA ────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("REPORTE EJECUTIVO CONTABLE")
    run.bold           = True
    run.font.size      = Pt(24)
    run.font.color.rgb = _rgb(COLOR_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SUCESION")
    run.bold           = True
    run.font.size      = Pt(16)
    run.font.color.rgb = _rgb("#2E86C1")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Periodo: {periodo}\nFecha de generacion: {datetime.now().strftime('%d de %B de %Y')}")
    run.font.size      = Pt(11)
    run.font.color.rgb = _rgb("#555555")

    doc.add_page_break()

    # ── Columnas especiales ────────────────────────────────────────────────────
    inquilino_col = next((c for c in tmp.columns if "inquilino" in c.lower()), None)
    inmueble_col  = next((c for c in tmp.columns if "inmueble"  in c.lower()), None)

    # ── Helpers para secciones de categoría ───────────────────────────────────
    def filter_by_kw(keywords):
        mask = pd.Series(False, index=tmp.index)
        for col in [conc_col, reg_col]:
            if col:
                s = tmp[col].astype(str).str.lower()
                for kw in keywords:
                    mask |= s.str.contains(kw.lower(), na=False, regex=False)
        return tmp[mask].copy()

    def tabla_transacciones(df_cat):
        if df_cat.empty:
            doc.add_paragraph("Sin registros para esta categoria.")
            return
        cols_hdr = ["Año", "Mes"]
        if conc_col: cols_hdr.append("Concepto")
        if reg_col:  cols_hdr.append("Registro")
        cols_hdr += ["Ingreso", "Egreso"]

        df_sorted = df_cat.sort_values(["_año", "_mes"], na_position="last")
        rows = []
        for _, r in df_sorted.iterrows():
            row = [
                str(int(r["_año"])) if pd.notna(r["_año"]) else "",
                MESES_CORTO[int(r["_mes"])-1] if pd.notna(r["_mes"]) and 1 <= int(r["_mes"]) <= 12 else "",
            ]
            if conc_col: row.append(str(r[conc_col])[:60])
            if reg_col:  row.append(str(r[reg_col])[:30])
            row += [
                money(r["_ing"]) if r["_ing"] > 0 else "",
                money(r["_egr"]) if r["_egr"] > 0 else "",
            ]
            rows.append(row)
        total_row = ["TOTAL"] + [""] * (len(cols_hdr) - 3) + [
            money(df_cat["_ing"].sum()), money(df_cat["_egr"].sum())
        ]
        rows.append(total_row)
        add_data_table(doc, cols_hdr, rows)

    def tabla_control(df_cat, col_cfg, importe_col):
        """col_cfg: [(col_df_o_None, encabezado), ...]  importe_col: '_ing' | '_egr'"""
        active = [(c, h) for c, h in col_cfg if c is not None]
        if df_cat.empty:
            doc.add_paragraph("Sin registros para esta categoria.")
            return
        hdrs = [h for _, h in active] + ["Importe"]
        df_s = df_cat.sort_values(["_año", "_mes"], na_position="last")
        rows = []
        for _, r in df_s.iterrows():
            row = []
            for col_key, _ in active:
                if col_key == "_año":
                    row.append(str(int(r["_año"])) if pd.notna(r["_año"]) else "")
                elif col_key == "_mes":
                    row.append(MESES_CORTO[int(r["_mes"])-1] if pd.notna(r["_mes"]) and 1 <= int(r["_mes"]) <= 12 else "")
                else:
                    val = r.get(col_key, "")
                    row.append(str(val)[:50] if pd.notna(val) else "")
            imp = r[importe_col]
            row.append(money(imp) if imp else "")
            rows.append(row)
        total = df_cat[importe_col].sum()
        rows.append(["TOTAL"] + [""] * (len(hdrs) - 2) + [money(total)])
        add_data_table(doc, hdrs, rows)

    # ── 1. RESUMEN EJECUTIVO ───────────────────────────────────────────────────
    add_heading(doc, "1. Resumen Ejecutivo", level=1)
    doc.add_paragraph(
        f"El presente reporte contable comprende el periodo {periodo} con un total de "
        f"{len(tmp):,} registros de movimientos financieros. "
        f"A continuacion se presentan los indicadores financieros clave."
    )
    doc.add_paragraph()

    add_kpi_table(doc, [
        ("Total Ingresos",  money(total_ing)),
        ("Total Egresos",   money(total_egr)),
        ("No. Registros",   f"{len(tmp):,}"),
        ("Periodo",         periodo),
    ])
    doc.add_paragraph()

    # Tabla anual de resumen
    add_heading(doc, "Resumen por Año", level=2)
    hdr_anual = ["Año", "Ingresos", "Egresos"]
    rows_anual = []
    for _, row in anual.iterrows():
        rows_anual.append([
            str(int(row["año"])),
            money(row["ing"]),
            money(row["egr"]),
        ])
    rows_anual.append(["TOTAL", money(anual["ing"].sum()), money(anual["egr"].sum())])
    add_data_table(doc, hdr_anual, rows_anual, col_widths=[2.5, 4.5, 4.5])
    doc.add_paragraph()

    # Grafica anual
    add_image(doc, chart_anual_barras(anual), width=Inches(6))
    doc.add_paragraph()
    doc.add_page_break()

    # ── 2. ANALISIS DE INGRESOS ────────────────────────────────────────────────
    add_heading(doc, "2. Analisis de Ingresos", level=1)
    doc.add_paragraph(
        f"Durante el periodo {periodo} se registraron ingresos totales por "
        f"{money(total_ing)}, distribuidos de la siguiente manera:"
    )

    # Por registro
    if reg_df is not None:
        add_heading(doc, "Ingresos por Registro", level=2)
        hdr_reg = [reg_col, "Ingresos", "% del Total"]
        rows_reg = []
        for _, row in reg_df.iterrows():
            pct = (row["ing"] / total_ing * 100) if total_ing else 0
            rows_reg.append([str(row[reg_col]), money(row["ing"]), f"{pct:.1f}%"])
        rows_reg.append(["TOTAL", money(total_ing), "100%"])
        add_data_table(doc, hdr_reg, rows_reg, col_widths=[5, 4, 3])
        doc.add_paragraph()

        add_image(doc, chart_pie_registro(reg_df, reg_col), width=Inches(5))
        doc.add_paragraph()

    add_heading(doc, "Control de Rentas", level=2)
    insertar_hoja_control(doc, "Control Rentas")
    doc.add_paragraph()

    # Mensual filtrado por RENTA (mismo universo que Control de Rentas)
    add_heading(doc, "Ingresos Mensuales — Rentas", level=2)
    if reg_col:
        _mask_renta = tmp[reg_col].astype(str).str.strip().str.lower() == "renta"
        if _mask_renta.sum() == 0:
            _mask_renta = tmp[reg_col].astype(str).str.strip().str.lower().str.contains("renta", regex=False)
        tmp_renta = tmp[_mask_renta]
    else:
        tmp_renta = tmp
    hdr_m = ["Mes"] + [str(int(a)) for a in años] + ["Total"]
    año_acum_ing = {a: 0.0 for a in años}
    rows_m = []
    for m in range(1, 13):
        row_vals = [MESES_LARGO[m-1]]
        row_total = 0.0
        for a in años:
            v = float(tmp_renta[(tmp_renta["_año"]==a) & (tmp_renta["_mes"]==m)]["_ing"].sum())
            row_vals.append(money(v))
            row_total += v
            año_acum_ing[a] += v
        row_vals.append(money(row_total))
        rows_m.append(row_vals)
    gran_total_ing = sum(año_acum_ing.values())
    col_totals = [money(año_acum_ing[a]) for a in años]
    rows_m.append(["TOTAL"] + col_totals + [money(gran_total_ing)])
    add_data_table(doc, hdr_m, rows_m)
    doc.add_paragraph()
    doc.add_page_break()

    # ── 3. ANALISIS DE EGRESOS ─────────────────────────────────────────────────
    add_heading(doc, "3. Analisis de Egresos", level=1)
    doc.add_paragraph(
        f"Los egresos totales del periodo ascienden a {money(total_egr)}."
    )

    if reg_df is not None:
        add_heading(doc, "Egresos por Registro", level=2)
        hdr_egr = [reg_col, "Egresos", "% del Total"]
        rows_egr = []
        for _, row in reg_df.iterrows():
            pct = (row["egr"] / total_egr * 100) if total_egr else 0
            rows_egr.append([str(row[reg_col]), money(row["egr"]), f"{pct:.1f}%"])
        rows_egr.append(["TOTAL", money(total_egr), "100%"])
        add_data_table(doc, hdr_egr, rows_egr, col_widths=[5, 4, 3])
        doc.add_paragraph()

    # Egresos filtrados a los mismos tipos que tienen Control (luz, predial, agua, despacho, impuestos)
    add_heading(doc, "Egresos Mensuales — Control de Egresos", level=2)
    _egr_tipos = ["luz", "predial", "agua", "honorarios despacho", "impuestos"]
    if reg_col:
        _mask_egr = pd.Series(False, index=tmp.index)
        for _tipo in _egr_tipos:
            _m = tmp[reg_col].astype(str).str.strip().str.lower() == _tipo.lower()
            if _m.sum() == 0:
                _m = tmp[reg_col].astype(str).str.strip().str.lower().str.contains(_tipo.lower(), regex=False)
            _mask_egr |= _m
        tmp_egr_ctrl = tmp[_mask_egr]
    else:
        tmp_egr_ctrl = tmp
    año_acum_egr = {a: 0.0 for a in años}
    rows_egr_m = []
    for m in range(1, 13):
        row_vals = [MESES_LARGO[m-1]]
        row_total = 0.0
        for a in años:
            v = float(tmp_egr_ctrl[(tmp_egr_ctrl["_año"]==a) & (tmp_egr_ctrl["_mes"]==m)]["_egr"].sum())
            row_vals.append(money(v))
            row_total += v
            año_acum_egr[a] += v
        row_vals.append(money(row_total))
        rows_egr_m.append(row_vals)
    gran_total_egr = sum(año_acum_egr.values())
    col_egr = [money(año_acum_egr[a]) for a in años]
    rows_egr_m.append(["TOTAL"] + col_egr + [money(gran_total_egr)])
    add_data_table(doc, hdr_m, rows_egr_m)
    doc.add_paragraph()

    for _titulo, _sheet in [
        ("Control de Luz",       "Control Luz"),
        ("Control de Predial",   "Control Predial"),
        ("Control de Agua",      "Control Agua"),
        ("Control de Despacho",  "Control Despacho"),
        ("Control de Impuestos", "Control Impuestos"),
    ]:
        add_heading(doc, _titulo, level=2)
        insertar_hoja_control(doc, _sheet)
        doc.add_paragraph()

    doc.add_page_break()

    # ── 4. REPORTE POR AÑO ────────────────────────────────────────────────────
    add_heading(doc, "4. Reporte Detallado por Año", level=1)

    for año in años:
        add_heading(doc, f"Año {año}", level=2)
        año_data = tmp[tmp["_año"] == año]
        ing_a = año_data["_ing"].sum()
        egr_a = año_data["_egr"].sum()
        bal_a = ing_a - egr_a

        add_kpi_table(doc, [
            ("Ingresos",  money(ing_a)),
            ("Egresos",   money(egr_a)),
            ("Registros", str(len(año_data))),
        ])
        doc.add_paragraph()

        # Tabla mensual del año
        hdr_año = ["Mes", "Ingresos", "Egresos"]
        rows_año = []
        for m in range(1, 13):
            md = año_data[año_data["_mes"] == m]
            if md.empty:
                continue
            i = md["_ing"].sum()
            e = md["_egr"].sum()
            rows_año.append([MESES_LARGO[m-1], money(i), money(e)])
        rows_año.append(["TOTAL", money(ing_a), money(egr_a)])
        add_data_table(doc, hdr_año, rows_año, col_widths=[5, 5.5, 5.5])
        doc.add_paragraph()

        # Grafica del año
        m_df = monthly_df(año_data)
        add_image(doc, chart_barras_mensual(m_df, año), width=Inches(6))

        if año != años[-1]:
            doc.add_page_break()

    doc.add_page_break()

    # ── 5. RESUMEN DETALLADO DE REGISTROS ─────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "5. Resumen Detallado de Registros", level=1)
    doc.add_paragraph(
        "Detalle completo de movimientos clasificados por tipo de registro, "
        "inmueble, año y mes, tomado de la hoja 'Detalle Registro' del dashboard."
    )
    doc.add_paragraph()
    insertar_detalle_registro(doc)
    doc.add_page_break()

    # ── 6. ANALISIS DE SITUACION DE RENTAS ────────────────────────────────────
    add_heading(doc, "6. Análisis de Situación de Rentas", level=1)

    doc.add_paragraph(
        "A continuación se describen las situaciones particulares que afectan "
        "el flujo de ingresos por concepto de renta durante el periodo reportado."
    )
    doc.add_paragraph()

    # — Caso 1: Jesús Emanuel Hurtado Murillo
    add_heading(doc, "Inquilino que dejó de rentar", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Jesús Emanuel Hurtado Murillo")
    run.bold = True
    p.add_run(
        " dejó de rentar en diciembre de 2025, por lo que no se registran "
        "ingresos por su concepto a partir del ejercicio 2026."
    )
    doc.add_paragraph()

    # — Caso 2: Rentas pendientes por mantenimiento
    add_heading(doc, "Rentas pendientes por acuerdo de mantenimiento", level=2)
    doc.add_paragraph(
        "Los siguientes inquilinos no han realizado el pago de rentas correspondientes "
        "al ejercicio 2026 en virtud de un acuerdo establecido con la sucesión, "
        "en el cual el pago queda condicionado a la realización de los mantenimientos "
        "correspondientes a los inmuebles arrendados:"
    )
    doc.add_paragraph()

    casos = [
        ("Li Shoy Yuan",    "Toledo 2103-B y Toledo 2101-A"),
        ("Weizhuang Zhen",  "Toledo 2103-A"),
    ]
    hdr_casos = ["Inquilino", "Inmueble(s)", "Situación"]
    rows_casos = [
        [nombre, inmueble,
         "Rentas 2026 pendientes — mantenimiento de inmueble en proceso"]
        for nombre, inmueble in casos
    ]
    add_data_table(doc, hdr_casos, rows_casos, col_widths=[4.5, 4.0, 7.0])
    doc.add_paragraph()

    doc.add_paragraph(
        "Una vez concluidos los trabajos de mantenimiento acordados, los inquilinos "
        "procederán a regularizar los pagos correspondientes. Se recomienda documentar "
        "formalmente el acuerdo y los plazos de entrega de las obras para dar seguimiento "
        "al cobro."
    )
    doc.add_paragraph()
    doc.add_page_break()

    # ── Pie de pagina / nota final ─────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— Fin del Reporte —")
    run.font.size      = Pt(10)
    run.font.color.rgb = _rgb("#888888")
    run.italic = True

    doc.save(output)
    print(f"  Documento guardado: {output}")


# ── Main ───────────────────────────────────────────────────────────────────────
def main():
    print("=" * 60)
    print("  Reporte Ejecutivo Contable — Sucesion")
    print("=" * 60)

    print("\nDescargando datos...")
    df = download_movimientos()
    print(f"Datos: {len(df)} filas x {len(df.columns)} columnas")

    print("\nGenerando reporte Word...")
    generate_report(df, OUTPUT_FILE)

    print(f"\n{'='*60}")
    print(f"  Archivo generado: {OUTPUT_FILE}")
    print(f"{'='*60}")

    import subprocess
    subprocess.Popen(["start", "", OUTPUT_FILE], shell=True)


if __name__ == "__main__":
    main()
