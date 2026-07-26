#!/usr/bin/env python3
"""
Reporte de Bitacora por Inmueble — Sucesion
Genera reporte_propiedades.docx: un salto de pagina por cada inmueble en
renta, con su bitacora completa por categoria (Renta, Luz, Agua, Predial,
Mantenimiento, Impuestos, etc.) y el detalle de productos de cada factura
XML relacionada a sus movimientos.
"""

import os
import sys
from datetime import datetime

import pandas as pd

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Dependencias faltantes. Ejecuta:  pip install -r requirements.txt")
    sys.exit(1)

import dashboard as dash
import reporte_word as rw

OUTPUT_FILE = "reporte_propiedades.docx"

# Orden preferido de categorias dentro de cada inmueble; el resto se agrega
# despues, en orden alfabetico.
ORDEN_CATEGORIAS = [
    "renta", "luz", "agua", "predial", "mantenimiento",
    "honorarios despacho", "impuestos",
]


def _orden_categoria(nombre: str):
    n = nombre.strip().lower()
    try:
        idx = ORDEN_CATEGORIAS.index(n)
    except ValueError:
        idx = len(ORDEN_CATEGORIAS)
    return (idx, n)


def _factura_bloque(doc, etiqueta, factura, importe_mov):
    """Inserta el encabezado de una factura y la tabla de productos detallados."""
    detalles = [etiqueta] if etiqueta else []
    if factura.get("emisor"):
        detalles.append(factura["emisor"])
    if factura.get("folio"):
        detalles.append(f"Folio {factura['folio']}")
    if factura.get("fecha"):
        detalles.append(factura["fecha"].strftime("%d/%m/%Y"))

    p = doc.add_paragraph()
    run = p.add_run("Factura — " + "  |  ".join(detalles) if detalles else "Factura")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = rw._rgb(rw.COLOR_DARK)

    hdr = ["Descripcion", "Importe"]
    rows = []
    for item in factura.get("conceptos", []):
        importe = item.get("importe") or 0
        rows.append([str(item.get("descripcion", "")), rw.money(importe) if importe else ""])
    if factura.get("iva"):
        rows.append(["IVA", rw.money(factura["iva"])])
    if factura.get("descuento"):
        rows.append(["Descuento", f"-{rw.money(factura['descuento'])}"])
    rows.append(["TOTAL FACTURA", rw.money(factura.get("total") or importe_mov)])

    rw.add_data_table(doc, hdr, rows, col_widths=[11.0, 5.0])
    doc.add_paragraph()


def render_propiedad(doc, propiedad, prop_data, facturas_idx, conc_col):
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(propiedad.upper())
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = rw._rgb(rw.COLOR_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bitacora completa de movimientos")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = rw._rgb("#555555")
    doc.add_paragraph()

    total_ing = float(prop_data["_ing_n"].sum())
    total_egr = float(prop_data["_egr_n"].sum())
    rw.add_kpi_table(doc, [
        ("Ingresos totales", rw.money(total_ing)),
        ("Egresos totales",  rw.money(total_egr)),
        ("Balance",          rw.money(total_ing - total_egr)),
    ])
    doc.add_paragraph()

    categorias = sorted(prop_data["_reg"].unique(), key=_orden_categoria)
    hdr_mov = ["Ano", "Mes", "Ingreso", "Egreso", "Concepto"]

    for cat in categorias:
        cat_data = prop_data[prop_data["_reg"] == cat].sort_values(
            ["_año", "_mes"], na_position="last")

        rw.add_heading(doc, cat.upper(), level=2)

        rows = []
        cat_ing = cat_egr = 0.0
        facturas_encontradas = []

        for _, row in cat_data.iterrows():
            año_v = str(int(row["_año"])) if pd.notna(row.get("_año")) else ""
            m_idx = row.get("_mes")
            mes_v = (dash.MESES_LARGO[int(m_idx) - 1]
                     if pd.notna(m_idx) and 1 <= int(m_idx) <= 12 else "")
            ing_v = float(row["_ing_n"])
            egr_v = float(row["_egr_n"])
            conc_v = str(row[conc_col]) if conc_col and pd.notna(row.get(conc_col)) else ""

            rows.append([
                año_v, mes_v,
                rw.money(ing_v) if ing_v else "",
                rw.money(egr_v) if egr_v else "",
                conc_v,
            ])
            cat_ing += ing_v
            cat_egr += egr_v

            if pd.notna(row.get("_fecha")):
                factura = dash._buscar_factura(
                    facturas_idx, row["_fecha"], row.get("_año"),
                    abs(egr_v) or abs(ing_v),
                )
                if factura:
                    etiqueta = f"{mes_v} {año_v}".strip()
                    facturas_encontradas.append((etiqueta, factura, abs(egr_v) or abs(ing_v)))

        if not rows:
            doc.add_paragraph("Sin movimientos registrados para esta categoria.")
            doc.add_paragraph()
            continue

        rows.append(["SUBTOTAL", "", rw.money(cat_ing), rw.money(cat_egr), ""])
        rw.add_data_table(doc, hdr_mov, rows, col_widths=[1.6, 2.2, 3.0, 3.0, 6.2])
        doc.add_paragraph()

        if facturas_encontradas:
            rw.add_heading(doc, "Detalle de facturas y productos", level=3)
            for etiqueta, factura, importe_mov in facturas_encontradas:
                _factura_bloque(doc, etiqueta, factura, importe_mov)
            doc.add_paragraph()


def generate_report(df: pd.DataFrame, output: str):
    fecha_col, ing_col, egr_col, conc_col, reg_col, _, inm_col, _, _ = dash.detect_columns(df)

    if not inm_col:
        print("ERROR: no se encontro columna Inmueble/Propiedad en los movimientos.")
        sys.exit(1)

    data = df.copy()
    data["_reg"] = (data[reg_col].astype(str).str.strip()
                     if reg_col else pd.Series("Sin registro", index=data.index))
    data["_inm"] = data[inm_col].astype(str).str.strip()
    data["_ing_n"] = dash.clean_numeric(data[ing_col]).fillna(0) if ing_col else 0
    data["_egr_n"] = dash.clean_numeric(data[egr_col]).fillna(0) if egr_col else 0
    data["_fecha"] = dash.parse_dates(data[fecha_col]) if fecha_col else pd.NaT
    data = dash._expandir_inmuebles_combinados(data)

    propiedades = dash.listar_propiedades(data, reg_col, "_inm")
    if not propiedades:
        print("ERROR: no hay inmuebles con Registro='renta' ni en "
              "dash.INMUEBLES_SIN_RENTA.")
        sys.exit(1)

    print(f"  Propiedades encontradas: {len(propiedades)}")

    print("  Indexando facturas XML...")
    facturas_idx = dash._indexar_facturas_xml()

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Portada ────────────────────────────────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BITACORA DE INMUEBLES EN RENTA")
    run.bold           = True
    run.font.size      = Pt(24)
    run.font.color.rgb = rw._rgb(rw.COLOR_DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SUCESION")
    run.bold           = True
    run.font.size      = Pt(16)
    run.font.color.rgb = rw._rgb("#2E86C1")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"{len(propiedades)} inmuebles  |  "
        f"Generado: {datetime.now().strftime('%d/%m/%Y')}"
    )
    run.font.size      = Pt(11)
    run.font.color.rgb = rw._rgb("#555555")

    # ── Una seccion por propiedad, con salto de pagina ────────────────────
    for propiedad in propiedades:
        prop_data = data[data["_inm"] == propiedad]
        render_propiedad(doc, propiedad, prop_data, facturas_idx, conc_col)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— Fin del reporte —")
    run.italic         = True
    run.font.size      = Pt(10)
    run.font.color.rgb = rw._rgb("#888888")

    doc.save(output)
    print(f"  Documento guardado: {output}")


def main():
    print("=" * 60)
    print("  Reporte de Bitacora por Inmueble — Sucesion")
    print("=" * 60)

    print("\nDescargando movimientos...")
    df = dash.download_movimientos()
    df = dash.prepare_año_mes(df)
    print(f"Datos: {len(df)} filas x {len(df.columns)} columnas")

    print("\nGenerando reporte Word...")
    generate_report(df, OUTPUT_FILE)

    print(f"\n{'='*60}")
    print(f"  Archivo generado: {OUTPUT_FILE}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
